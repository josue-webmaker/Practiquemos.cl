from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

GREEN = RGBColor(0x1B, 0x8A, 0x2F)
DARK_GREEN = RGBColor(0x14, 0x6B, 0x23)
ORANGE = RGBColor(0xD9, 0x77, 0x06)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
BLUE = RGBColor(0x1D, 0x4E, 0xD8)
RED = RGBColor(0xDC, 0x26, 0x26)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

logo_path = 'attached_assets/logo_optimized.png'

def add_heading_styled(text, level=1, color=GREEN):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return heading

def add_thin_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:color'), '1B8A2F')
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def style_cell(cell, text, bold=False, color=DARK_GRAY, bg=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    if bg:
        set_cell_shading(cell, bg)

def add_bullet(title, desc, title_color=DARK_GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run_t = p.add_run(f'  {title}: ')
    run_t.bold = True
    run_t.font.color.rgb = title_color
    run_t.font.size = Pt(11)
    run_d = p.add_run(desc)
    run_d.font.color.rgb = DARK_GRAY
    run_d.font.size = Pt(11)

def add_check(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f'  {text}')
    run.font.color.rgb = DARK_GREEN
    run.font.size = Pt(11)

def add_warning(title, desc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run_t = p.add_run(f'  {title}: ')
    run_t.bold = True
    run_t.font.color.rgb = ORANGE
    run_t.font.size = Pt(11)
    run_d = p.add_run(desc)
    run_d.font.color.rgb = DARK_GRAY
    run_d.font.size = Pt(11)

def add_body(text, indent=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    for run in p.runs:
        run.font.color.rgb = DARK_GRAY
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    return p

def add_body_bold_start(bold_text, normal_text, bold_color=DARK_GREEN):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run_b = p.add_run(bold_text)
    run_b.bold = True
    run_b.font.color.rgb = bold_color
    run_b.font.size = Pt(11)
    run_n = p.add_run(normal_text)
    run_n.font.color.rgb = DARK_GRAY
    run_n.font.size = Pt(11)

def add_spacer(pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)

# ============================================================
#                       PORTADA
# ============================================================

add_spacer(40)

p_logo = doc.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(logo_path):
    run_logo = p_logo.add_run()
    run_logo.add_picture(logo_path, width=Inches(2.5))

add_spacer(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DOCUMENTO DE ENTREGA')
run.font.size = Pt(32)
run.font.color.rgb = GREEN
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROYECTO')
run.font.size = Pt(14)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Practiquemos.cl')
run.font.size = Pt(26)
run.font.color.rgb = DARK_GREEN
run.bold = True
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Plataforma de preparación para el examen\nde licencia de conducir en Chile')
run.font.size = Pt(13)
run.font.color.rgb = MEDIUM_GRAY
run.italic = True
run.font.name = 'Calibri'

add_spacer(30)

# Info box
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl.rows[0].cells[0]
set_cell_shading(cell, 'F0FFF4')
cell.width = Cm(12)
cell.text = ''
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after = Pt(12)

for line, color, bold in [
    ('Fecha de entrega: 18 de marzo de 2026', MEDIUM_GRAY, False),
    ('Versión: 1.0', MEDIUM_GRAY, False),
    ('Documento confidencial', LIGHT_GRAY, False),
    ('', MEDIUM_GRAY, False),
    ('Desarrollado por:', MEDIUM_GRAY, False),
    ('Lucas - WebMakerChile', GREEN, True),
    ('webmakerchile@gmail.com', MEDIUM_GRAY, False),
    ('+56 9 6251 1821', MEDIUM_GRAY, False),
]:
    if not line and not bold:
        run = p.add_run('\n')
        continue
    run = p.add_run(line + '\n')
    run.font.size = Pt(10)
    run.font.color.rgb = color
    run.bold = bold
    run.font.name = 'Calibri'

doc.add_page_break()

# ============================================================
#                       ÍNDICE
# ============================================================

add_heading_styled('Índice de contenidos', level=1)
add_thin_line()

indice = [
    ('1.', 'Carta de presentación'),
    ('2.', 'Resumen ejecutivo del proyecto'),
    ('3.', 'Proceso de desarrollo'),
    ('4.', 'Arquitectura y tecnologías utilizadas'),
    ('5.', 'Integraciones implementadas'),
    ('6.', 'Desafíos superados y soluciones'),
    ('7.', 'Cómo empezar a usar la plataforma'),
    ('8.', 'Funcionalidades completas'),
    ('9.', 'Panel de administrador'),
    ('10.', 'Gestión de preguntas'),
    ('11.', 'Planes, pagos y modelo de negocio'),
    ('12.', 'Configuración de exámenes por licencia'),
    ('13.', 'Seguridad y protección de datos'),
    ('14.', 'Estado de entrega y próximos pasos'),
    ('15.', 'Garantía y soporte técnico'),
]

for num, title in indice:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1)
    run_n = p.add_run(f'{num}  ')
    run_n.font.color.rgb = GREEN
    run_n.bold = True
    run_n.font.size = Pt(11)
    run_t = p.add_run(title)
    run_t.font.color.rgb = DARK_GRAY
    run_t.font.size = Pt(11)

doc.add_page_break()

# ============================================================
#           1. CARTA DE PRESENTACIÓN
# ============================================================

add_heading_styled('1. Carta de presentación', level=1)
add_thin_line()

add_body('Estimado cliente,')

add_body(
    'Es un honor para WebMakerChile hacer entrega formal del proyecto Practiquemos.cl, '
    'una plataforma integral de preparación para el examen de licencia de conducir en Chile. '
    'Este documento detalla el trabajo realizado, las decisiones tomadas durante el desarrollo, '
    'los desafíos que se presentaron y las soluciones implementadas.'
)

add_body(
    'Practiquemos.cl nació de una visión clara: crear una herramienta accesible, moderna y efectiva '
    'que acompañe a miles de chilenos en su camino hacia la obtención de la licencia de conducir. '
    'Desde el primer día, nos comprometimos a construir algo que no fuera solo funcional, sino que '
    'ofreciera una experiencia de usuario excepcional, con elementos de gamificación, inteligencia '
    'artificial y un diseño pensado para motivar al usuario en cada paso del estudio.'
)

add_body(
    'A lo largo de este documento encontrará una descripción completa de todo lo construido, '
    'las tecnologías empleadas, las integraciones realizadas y las instrucciones necesarias para '
    'que usted pueda administrar la plataforma de forma autónoma. Nuestro objetivo es que tenga '
    'total claridad y confianza sobre el producto que recibe.'
)

add_body(
    'Agradecemos la confianza depositada en nuestro equipo. Estamos orgullosos del resultado '
    'y comprometidos a seguir apoyándole en las etapas que vienen.'
)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
run = p.add_run('Atentamente,')
run.font.color.rgb = DARK_GRAY
run.italic = True

p = doc.add_paragraph()
run = p.add_run('Lucas')
run.font.color.rgb = GREEN
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
run = p.add_run('Fundador y Desarrollador Principal')
run.font.color.rgb = MEDIUM_GRAY
run.font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('WebMakerChile')
run.font.color.rgb = GREEN
run.bold = True
run.font.size = Pt(10)

doc.add_page_break()

# ============================================================
#           2. RESUMEN EJECUTIVO
# ============================================================

add_heading_styled('2. Resumen ejecutivo del proyecto', level=1)
add_thin_line()

add_body(
    'Practiquemos.cl es una aplicación multiplataforma (web, iOS y Android) diseñada para '
    'preparar a los usuarios para el examen teórico de licencia de conducir en Chile. '
    'La plataforma cubre los 6 tipos de licencia existentes y ofrece una experiencia de '
    'estudio completa, interactiva y gamificada.'
)

# Key metrics table
tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

metrics = [
    ('Cifras clave del proyecto', ''),
    ('Total de preguntas', '1.591 preguntas (306 oficiales de CONASET)'),
    ('Tipos de licencia', '6 (Clase B, A2, A4, C, D y E)'),
    ('Plataformas', 'Web, iOS y Android desde un solo código fuente'),
]

style_cell(tbl.rows[0].cells[0], metrics[0][0], bold=True, color=WHITE, bg='1B8A2F', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
tbl.rows[0].cells[0].merge(tbl.rows[0].cells[1])

for i, (label, value) in enumerate(metrics[1:], 1):
    style_cell(tbl.rows[i].cells[0], label, bold=True, color=DARK_GREEN, bg='F0FFF4', size=10)
    style_cell(tbl.rows[i].cells[1], value, color=DARK_GRAY, size=10)

add_spacer()

add_heading_styled('Propuesta de valor', level=2, color=DARK_GREEN)

value_props = [
    ('Contenido oficial', 'Incluye preguntas basadas en el material oficial de CONASET, '
     'garantizando relevancia y actualidad.'),
    ('Experiencia gamificada', 'Una mascota copiloto animada acompaña al usuario, generando '
     'motivación a través de refuerzos positivos y rachas de aciertos.'),
    ('Inteligencia artificial', 'Integración con OpenAI para lectura en voz alta con voz natural, '
     'haciendo el contenido accesible para todos.'),
    ('Modelo freemium', 'Acceso gratuito limitado con opción de planes premium, permitiendo '
     'monetización sin barreras de entrada.'),
    ('Multiplataforma', 'Una sola base de código para web, iOS y Android, optimizando costos '
     'de mantenimiento y actualizaciones.'),
]

for title, desc in value_props:
    add_bullet(title, desc)

doc.add_page_break()

# ============================================================
#           3. PROCESO DE DESARROLLO
# ============================================================

add_heading_styled('3. Proceso de desarrollo', level=1)
add_thin_line()

add_body(
    'El desarrollo de Practiquemos.cl fue un proceso riguroso que abarcó múltiples etapas, '
    'desde la investigación inicial del contenido oficial hasta la implementación de integraciones '
    'complejas con servicios de pago y distribución en tiendas de aplicaciones.'
)

add_heading_styled('Fase 1: Investigación y planificación', level=2, color=DARK_GREEN)
add_body(
    'Se realizó un análisis exhaustivo del examen de licencia de conducir en Chile, revisando '
    'el material oficial de CONASET para cada uno de los 6 tipos de licencia. Se definieron '
    'los parámetros exactos de cada examen: cantidad de preguntas, puntajes de aprobación, '
    'tiempos límite y categorías temáticas. Esta fase fue fundamental para garantizar que '
    'la plataforma reflejara con precisión la experiencia del examen real.'
)

add_heading_styled('Fase 2: Diseño de la experiencia de usuario', level=2, color=DARK_GREEN)
add_body(
    'Se diseñó una interfaz centrada en el usuario, con una paleta de colores profesional '
    '(azul primario, ámbar como acento y tonos neutros), tipografía Nunito para máxima legibilidad '
    'y un sistema de retroalimentación visual que evita el uso de colores agresivos para los errores, '
    'optando por tonos anaranjados más amigables. El concepto de la mascota copiloto se desarrolló '
    'como elemento central de la experiencia, con 5 estados visuales diferentes que responden '
    'a las acciones del usuario.'
)

add_heading_styled('Fase 3: Desarrollo del motor de exámenes', level=2, color=DARK_GREEN)
add_body(
    'Se construyó un motor de exámenes robusto y configurable que soporta las particularidades '
    'de cada tipo de licencia. Incluye temporizador regresivo, navegación entre preguntas, '
    'barajado aleatorio de opciones (para evitar memorización por posición), modo de aprendizaje '
    'con explicaciones, y un sistema de puntuación que contempla incluso preguntas con puntaje '
    'doble en la Clase B. Se implementaron efectos de sonido con 3 variaciones para respuestas '
    'correctas e incorrectas, seleccionadas aleatoriamente para mantener la frescura.'
)

add_heading_styled('Fase 4: Base de datos y contenido', level=2, color=DARK_GREEN)
add_body(
    'Se estructuró una base de datos PostgreSQL con más de 1.591 preguntas distribuidas en '
    '10 categorías temáticas. Cada pregunta incluye: enunciado, 4 opciones de respuesta, '
    'respuesta correcta, explicación detallada, nivel de dificultad y tipos de licencia aplicables. '
    'Se incorporaron 306 preguntas oficiales de CONASET y se creó un sistema de imágenes con '
    '68 ilustraciones que se asignan automáticamente a las preguntas mediante un motor de '
    'coincidencia por palabras clave, logrando una cobertura visual del 100%.'
)

add_heading_styled('Fase 5: Temario y material de estudio', level=2, color=DARK_GREEN)
add_body(
    'Se desarrolló un temario completo con 11 capítulos: 6 capítulos base aplicables a todos '
    'los tipos de licencia (con 83 secciones) y 5 capítulos específicos por categoría (motocicletas, '
    'transporte de pasajeros, transporte de carga, maquinaria pesada y tracción animal). '
    'Cada sección incluye contenido detallado y utiliza las mismas imágenes ilustrativas '
    'del motor de exámenes.'
)

add_heading_styled('Fase 6: Integraciones y servicios', level=2, color=DARK_GREEN)
add_body(
    'Se implementaron múltiples integraciones con servicios externos: sistema de pagos con '
    'Mercado Pago para web y Android, RevenueCat para compras in-app en iOS, y OpenAI para '
    'la función de lectura en voz alta con inteligencia artificial. Cada integración requirió '
    'su propia configuración de seguridad, manejo de errores y flujos de activación.'
)

add_heading_styled('Fase 7: Panel de administración', level=2, color=DARK_GREEN)
add_body(
    'Se construyó un panel de administración completo que permite la gestión integral '
    'de usuarios y contenido: creación, edición y eliminación de usuarios con confirmación '
    'de seguridad, búsqueda y filtrado avanzado, gestión de planes premium, y un módulo '
    'completo de administración de preguntas con búsqueda, filtros, creación, edición y '
    'desactivación de contenido.'
)

add_heading_styled('Fase 8: Cumplimiento y compliance', level=2, color=DARK_GREEN)
add_body(
    'Se implementaron todas las medidas necesarias para cumplir con los requisitos de '
    'Apple App Store y Google Play Store: política de privacidad completa, términos de servicio, '
    'declaración de uso de encriptación, manifiestos de privacidad, botón de restaurar compras '
    'en iOS, gestión de reembolsos y toda la documentación legal requerida.'
)

doc.add_page_break()

# ============================================================
#           4. ARQUITECTURA Y TECNOLOGÍAS
# ============================================================

add_heading_styled('4. Arquitectura y tecnologías utilizadas', level=1)
add_thin_line()

add_body(
    'La plataforma fue construida con una arquitectura moderna que prioriza el rendimiento, '
    'la escalabilidad y la mantenibilidad a largo plazo.'
)

add_heading_styled('Aplicación móvil y web (Frontend)', level=2, color=DARK_GREEN)

tech_front = [
    ('Expo SDK 54 con React Native', 'Framework que permite desarrollar una sola aplicación '
     'que funciona en web, iOS y Android simultáneamente.'),
    ('TypeScript', 'Lenguaje con tipado estático que reduce errores y mejora la calidad del código.'),
    ('React Navigation (expo-router)', 'Sistema de navegación basado en archivos, similar '
     'a los estándares modernos de la industria.'),
    ('TanStack React Query', 'Gestión inteligente de datos del servidor con caché automático '
     'y sincronización en tiempo real.'),
    ('React Native Reanimated', 'Motor de animaciones de alto rendimiento para la mascota '
     'copiloto y las transiciones de la interfaz.'),
]

for title, desc in tech_front:
    add_bullet(title, desc)

add_heading_styled('Servidor (Backend)', level=2, color=DARK_GREEN)

tech_back = [
    ('Node.js con Express 5', 'Servidor web robusto y eficiente para manejar las solicitudes '
     'de la aplicación.'),
    ('TypeScript', 'El mismo lenguaje en frontend y backend, facilitando el mantenimiento.'),
    ('PostgreSQL', 'Base de datos relacional de nivel empresarial para almacenar toda la '
     'información de forma segura.'),
    ('Drizzle ORM', 'Interfaz moderna de acceso a datos con validación de esquemas y migraciones '
     'automáticas.'),
    ('bcrypt', 'Algoritmo de encriptación de contraseñas de grado militar.'),
]

for title, desc in tech_back:
    add_bullet(title, desc)

doc.add_page_break()

# ============================================================
#           5. INTEGRACIONES
# ============================================================

add_heading_styled('5. Integraciones implementadas', level=1)
add_thin_line()

add_heading_styled('Mercado Pago (Checkout Pro)', level=2, color=DARK_GREEN)
add_body(
    'Se integró Mercado Pago como pasarela de pagos principal para usuarios de web y Android. '
    'El sistema crea preferencias de pago con los datos del plan seleccionado, redirige al '
    'usuario al checkout seguro de Mercado Pago, y al completarse el pago, un webhook notifica '
    'al servidor para activar automáticamente el plan premium del usuario. Se implementó '
    'verificación de firmas para garantizar la autenticidad de las notificaciones.'
)

add_heading_styled('RevenueCat (compras in-app iOS)', level=2, color=DARK_GREEN)
add_body(
    'Para cumplir con las políticas de Apple, se integró RevenueCat como intermediario '
    'para las compras dentro de la aplicación en iOS. RevenueCat gestiona la comunicación '
    'con App Store Connect, valida las compras y notifica al servidor mediante webhooks. '
    'Se implementó un botón de "Restaurar compras" en la pantalla de planes (visible solo en iOS) '
    'como requisito obligatorio de Apple. La verificación de planes se realiza del lado del servidor '
    'para prevenir manipulaciones del cliente.'
)

add_heading_styled('OpenAI TTS HD (lectura en voz alta)', level=2, color=DARK_GREEN)
add_body(
    'Se integró el servicio de text-to-speech de OpenAI con la voz "Nova" (voz femenina, '
    'profesional y clara) para leer en voz alta las preguntas, explicaciones y secciones del '
    'temario. El texto se procesa previamente para agregar pausas naturales antes de enviarlo '
    'a la API. El audio se genera en el servidor y se reproduce en el cliente, con botones '
    'de altavoz disponibles en todas las pantallas relevantes. Se incluye una pantalla de '
    'configuración de voz donde el usuario puede previsualizar el sonido.'
)

add_heading_styled('Expo AV (efectos de sonido)', level=2, color=DARK_GREEN)
add_body(
    'Se implementó un sistema de efectos de sonido con archivos WAV locales para las respuestas '
    'correctas e incorrectas. Se crearon 3 variaciones para cada tipo de respuesta, y el sistema '
    'selecciona aleatoriamente una variación diferente cada vez para mantener la experiencia fresca '
    'y evitar la monotonía.'
)

doc.add_page_break()

# ============================================================
#           6. DESAFÍOS Y SOLUCIONES
# ============================================================

add_heading_styled('6. Desafíos superados y soluciones', level=1)
add_thin_line()

add_body(
    'Durante el desarrollo se presentaron diversos desafíos técnicos y de diseño que '
    'requirieron soluciones creativas e innovadoras:'
)

challenges = [
    ('Sesgo en las respuestas correctas',
     'Se detectó que en los datos originales, el 83% de las respuestas correctas estaban '
     'en la opción B. Esto permitiría a los usuarios aprobar simplemente seleccionando siempre '
     'la opción B. Se implementó un algoritmo de barajado que redistribuye aleatoriamente '
     'las opciones al inicio de cada examen, eliminando completamente este sesgo.',
     'Barajado aleatorio de opciones en cada examen.'),

    ('Compatibilidad multiplataforma',
     'Desarrollar una aplicación que funcione correctamente en web, iOS y Android '
     'simultáneamente presentó desafíos significativos, especialmente en el manejo de '
     'audio, animaciones y gestos táctiles. Se utilizaron abstracciones de plataforma y '
     'detección de sistema operativo para adaptar el comportamiento en cada entorno.',
     'Código adaptativo con detección de plataforma.'),

    ('Cumplimiento de Apple App Store',
     'Las políticas de Apple requieren múltiples elementos específicos: botón de restaurar '
     'compras, declaración de encriptación, manifiestos de privacidad, política de privacidad '
     'accesible por URL, y uso de su sistema de compras para transacciones digitales. Se '
     'implementaron todos estos requisitos de forma exhaustiva.',
     'Implementación completa de todos los requisitos de compliance.'),

    ('Seguridad en la verificación de compras',
     'Inicialmente, la activación de planes premium confiaba en datos enviados por el cliente, '
     'lo que podría permitir que un usuario malintencionado se otorgara un plan superior al '
     'adquirido. Se rediseñó el sistema para que la verificación se realice completamente '
     'del lado del servidor, consultando directamente a RevenueCat los derechos reales del usuario.',
     'Verificación de planes del lado del servidor.'),

    ('Cobertura visual de preguntas',
     'Con 1.591 preguntas y 68 imágenes disponibles, el desafío era asignar imágenes relevantes '
     'a cada pregunta. Se desarrolló un motor inteligente de coincidencia por palabras clave '
     'que analiza el contenido de cada pregunta y selecciona la imagen más apropiada, logrando '
     'una cobertura del 100% sin necesidad de asignación manual.',
     'Motor de coincidencia por palabras clave con 100% de cobertura.'),

    ('Migración de contraseñas',
     'El sistema originalmente usaba un algoritmo de encriptación básico (SHA-256). Se migró '
     'a bcrypt (estándar de la industria) con un sistema de migración automática: cuando un '
     'usuario con contraseña antigua inicia sesión, su contraseña se actualiza automáticamente '
     'al nuevo formato sin que el usuario lo note.',
     'Migración transparente y automática al iniciar sesión.'),
]

for title, desc, solution in challenges:
    add_heading_styled(title, level=3, color=DARK_GREEN)
    add_body(desc)
    add_body_bold_start('Solución: ', solution)
    add_spacer(4)

doc.add_page_break()

# ============================================================
#           7. CÓMO EMPEZAR
# ============================================================

add_heading_styled('7. Cómo empezar a usar la plataforma', level=1)
add_thin_line()

add_heading_styled('Paso 1: Crear su cuenta', level=2, color=DARK_GREEN)
add_body('Para comenzar a administrar la plataforma, primero debe registrarse como usuario:')

steps = [
    'Ingrese a la aplicación desde su navegador web o dispositivo móvil.',
    'Toque el botón "Registrarse" en la pantalla de inicio.',
    'Complete el formulario con sus datos: nombre de usuario, contraseña, nombre completo y correo electrónico.',
    'Una vez registrado, comunique su nombre de usuario al equipo de WebMakerChile a través de los canales de contacto indicados al final de este documento.',
]

for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f'    {i}.  {step}')
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.size = Pt(11)

add_heading_styled('Paso 2: Activación de permisos de administrador', level=2, color=DARK_GREEN)
add_body(
    'Una vez que WebMakerChile reciba su nombre de usuario, se le otorgarán los permisos '
    'de administrador desde el servidor. A partir de ese momento, usted tendrá acceso completo '
    'al Panel de Administrador desde el menú principal de la aplicación.'
)

# Warning box
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl.rows[0].cells[0]
set_cell_shading(cell, 'FEF2F2')
cell.text = ''
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(10)
run = p.add_run('    Importante: ')
run.bold = True
run.font.color.rgb = RED
run.font.size = Pt(11)
run = p.add_run(
    'No comparta sus credenciales de acceso con terceros. Usted es el responsable '
    'de la seguridad de su cuenta de administrador.'
)
run.font.color.rgb = DARK_GRAY
run.font.size = Pt(11)

doc.add_page_break()

# ============================================================
#           8. FUNCIONALIDADES COMPLETAS
# ============================================================

add_heading_styled('8. Funcionalidades completas', level=1)
add_thin_line()

add_body(
    'A continuación se detallan todas las funcionalidades disponibles para los usuarios '
    'de la plataforma:'
)

features = [
    ('Soporte para 6 tipos de licencia',
     'Clase B (vehículos particulares), A2 (transporte de pasajeros), A4 (transporte de carga), '
     'C (motocicletas), D (maquinaria pesada) y E (tracción animal). Cada licencia tiene su '
     'propio conjunto de preguntas, configuración de examen y capítulos de estudio específicos.'),
    ('1.591 preguntas con explicaciones',
     'Cada pregunta incluye 4 alternativas, la respuesta correcta, una explicación detallada '
     'y su categoría temática. Las 306 preguntas oficiales de CONASET están identificadas '
     'con una marca especial.'),
    ('Dos modos de examen',
     'El modo Simulación replica las condiciones del examen real con temporizador y puntaje. '
     'El modo Aprendizaje muestra la explicación de cada pregunta inmediatamente después de '
     'responder, permitiendo un estudio más reflexivo.'),
    ('Temario de estudio',
     '11 capítulos organizados temáticamente con un total de 115 secciones. Los 6 capítulos '
     'base son accesibles para todas las licencias, y los 5 capítulos específicos se muestran '
     'según el tipo de licencia seleccionado por el usuario.'),
    ('Sistema de progreso',
     'Seguimiento detallado del avance por categoría temática. El usuario puede ver qué '
     'temas domina y cuáles necesitan más práctica, con indicadores visuales de porcentaje.'),
    ('Historial de exámenes',
     'Registro completo de todos los exámenes realizados, con fecha, puntaje obtenido, '
     'resultado (aprobado/reprobado), tiempo empleado y desglose por categoría.'),
    ('Preguntas favoritas',
     'Los usuarios pueden marcar preguntas para repasarlas después, creando una lista '
     'personalizada de estudio enfocado.'),
    ('Lectura en voz alta con IA',
     'Todas las preguntas, explicaciones y secciones del temario pueden escucharse con '
     'voz natural generada por inteligencia artificial (OpenAI Nova). Incluye pantalla '
     'de configuración de voz.'),
    ('Mascota copiloto animada',
     'Un compañero visual que cambia de estado según las acciones del usuario: celebra los '
     'aciertos, piensa durante las pausas, habla cuando se activa la lectura en voz alta '
     'y motiva con animaciones positivas. Usa 5 imágenes reales con transiciones fluidas.'),
    ('Efectos de sonido',
     '3 variaciones de sonido para respuestas correctas y 3 para incorrectas, seleccionadas '
     'aleatoriamente para evitar monotonía.'),
    ('68 imágenes ilustrativas',
     'Motor inteligente de asignación de imágenes por palabras clave. Cada pregunta muestra '
     'una imagen relevante que ilustra la situación de tránsito descrita.'),
    ('Insignias por rachas',
     'Sistema de gamificación que premia las rachas de respuestas correctas consecutivas, '
     'incentivando la práctica constante.'),
]

for title, desc in features:
    add_bullet(title, desc)

doc.add_page_break()

# ============================================================
#           9. PANEL DE ADMINISTRADOR
# ============================================================

add_heading_styled('9. Panel de administrador', level=1)
add_thin_line()

add_body(
    'El panel de administración le otorga control total sobre la plataforma. Está diseñado '
    'para ser intuitivo y completo, permitiendo gestionar todos los aspectos del negocio '
    'sin necesidad de conocimientos técnicos.'
)

add_heading_styled('Gestión de usuarios', level=2, color=DARK_GREEN)

admin_users = [
    ('Crear usuarios',
     'Registre nuevos usuarios manualmente con todos sus datos: nombre de usuario, contraseña, '
     'nombre completo, correo electrónico, tipo de licencia, plan y rol.'),
    ('Editar usuarios',
     'Modifique cualquier dato de un usuario existente: nombre completo, correo electrónico, '
     'contraseña, plan (gratuito, premium 10 días o premium 30 días) y rol '
     '(usuario estándar o administrador).'),
    ('Eliminar usuarios',
     'Elimine usuarios del sistema. Se incluye un diálogo de confirmación con advertencia '
     'para prevenir eliminaciones accidentales.'),
    ('Buscar usuarios',
     'Barra de búsqueda en tiempo real que filtra por nombre, nombre de usuario o correo electrónico.'),
    ('Filtrar por categoría',
     'Las tarjetas de estadísticas son interactivas: al tocar "Premium" se muestran solo '
     'los usuarios premium, al tocar "Gratuitos" solo los gratuitos, y así sucesivamente.'),
    ('Ver ficha detallada',
     'Al tocar cualquier usuario se abre una ficha completa con: nombre de usuario, nombre '
     'completo, correo electrónico, tipo de licencia, plan activo, vigencia del plan '
     '(con indicador de color: verde si quedan días, ámbar si está por vencer, rojo si expiró), '
     'fecha de registro y fecha del último inicio de sesión.'),
    ('Otorgar o revocar premium',
     'Cambie el plan de cualquier usuario directamente desde el panel. Al asignar un plan '
     'premium, la vigencia se calcula automáticamente (10 o 30 días desde la fecha de cambio).'),
]

for title, desc in admin_users:
    add_bullet(title, desc)

add_heading_styled('Estadísticas del panel', level=2, color=DARK_GREEN)
add_body(
    'El panel presenta 4 tarjetas de estadísticas en la parte superior: total de usuarios '
    'registrados, cantidad de usuarios premium activos, cantidad de usuarios con plan gratuito '
    'y cantidad de administradores. Estas tarjetas se actualizan en tiempo real y funcionan '
    'como filtros rápidos.'
)

doc.add_page_break()

# ============================================================
#           10. GESTIÓN DE PREGUNTAS
# ============================================================

add_heading_styled('10. Gestión de preguntas', level=1)
add_thin_line()

add_body(
    'El módulo de gestión de preguntas le permite mantener y actualizar el contenido '
    'de la plataforma de forma completa:'
)

question_features = [
    ('Búsqueda avanzada',
     'Busque preguntas por texto del enunciado, con filtros combinables por categoría '
     'temática, tipo de licencia, nivel de dificultad y origen (oficial CONASET o propia).'),
    ('Crear nuevas preguntas',
     'Agregue preguntas con: enunciado, 4 opciones de respuesta, indicación de la respuesta '
     'correcta, explicación detallada, categoría, nivel de dificultad y tipos de licencia aplicables.'),
    ('Editar preguntas existentes',
     'Modifique cualquier campo de las preguntas ya existentes en la base de datos.'),
    ('Desactivar preguntas',
     'En lugar de eliminar, las preguntas se desactivan (se ocultan). Esto permite reactivarlas '
     'en el futuro si es necesario, sin perder el contenido.'),
    ('Identificación de preguntas oficiales',
     'Las preguntas provenientes de CONASET están marcadas con una etiqueta especial que las '
     'diferencia del contenido propio de la plataforma.'),
    ('Paginación',
     'El listado de preguntas incluye paginación para navegar eficientemente entre las '
     '1.591 preguntas de la base de datos.'),
]

for title, desc in question_features:
    add_bullet(title, desc)

doc.add_page_break()

# ============================================================
#           11. PLANES, PAGOS Y MODELO DE NEGOCIO
# ============================================================

add_heading_styled('11. Planes, pagos y modelo de negocio', level=1)
add_thin_line()

add_heading_styled('Estructura de planes', level=2, color=DARK_GREEN)

add_body(
    'La plataforma opera con un modelo freemium que permite a los usuarios probar la '
    'aplicación de forma gratuita antes de decidir si desean acceder al contenido completo:'
)

tbl = doc.add_table(rows=4, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

for i, h in enumerate(['Plan', 'Precio', 'Duración', 'Acceso']):
    style_cell(tbl.rows[0].cells[i], h, bold=True, color=WHITE, bg='1B8A2F', align=WD_ALIGN_PARAGRAPH.CENTER)

plans = [
    ('Gratuito', 'Sin costo', 'Permanente', 'Limitado'),
    ('Premium 10 días', '$2.990 CLP', '10 días', 'Completo'),
    ('Premium 30 días', '$4.990 CLP', '30 días', 'Completo'),
]

for i, (plan, price, dur, access) in enumerate(plans, 1):
    style_cell(tbl.rows[i].cells[0], plan, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell(tbl.rows[i].cells[1], price, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell(tbl.rows[i].cells[2], dur, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell(tbl.rows[i].cells[3], access, align=WD_ALIGN_PARAGRAPH.CENTER)

add_spacer()

add_heading_styled('Métodos de pago integrados', level=2, color=DARK_GREEN)

add_body_bold_start('Mercado Pago (web y Android): ',
    'Los usuarios seleccionan su plan en la pantalla de planes, son redirigidos al checkout '
    'seguro de Mercado Pago, y al completar el pago, el plan se activa automáticamente '
    'mediante una notificación webhook al servidor.')

add_body_bold_start('RevenueCat (iOS): ',
    'En dispositivos Apple, las compras se procesan a través del sistema nativo de compras '
    'de la App Store, gestionado por RevenueCat. Incluye soporte para restaurar compras '
    'previas, como lo requiere Apple.')

add_body_bold_start('Asignación manual: ',
    'Usted como administrador puede otorgar planes premium directamente desde el panel '
    'de administración, sin que el usuario necesite realizar un pago. Útil para promociones, '
    'cortesías o soporte al cliente.')

doc.add_page_break()

# ============================================================
#           12. CONFIGURACIÓN DE EXÁMENES
# ============================================================

add_heading_styled('12. Configuración de exámenes por licencia', level=1)
add_thin_line()

add_body(
    'Cada tipo de licencia tiene una configuración específica basada en los parámetros '
    'oficiales del examen teórico en Chile:'
)

tbl = doc.add_table(rows=7, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

for i, h in enumerate(['Licencia', 'Preguntas', 'Puntaje mínimo', 'Tiempo']):
    style_cell(tbl.rows[0].cells[i], h, bold=True, color=WHITE, bg='1B8A2F', align=WD_ALIGN_PARAGRAPH.CENTER)

exams = [
    ('Clase B', '35', '33/38 (87%)', '45 minutos'),
    ('Clase A2', '20', '16/20 (80%)', '30 minutos'),
    ('Clase A4', '35', '70%', '45 minutos'),
    ('Clase C', '20', '15/20 (75%)', '30 minutos'),
    ('Clase D', '12', '9/12 (75%)', '20 minutos'),
    ('Clase E', '10', '7/10 (70%)', '20 minutos'),
]

for i, (lic, qs, score, time) in enumerate(exams, 1):
    style_cell(tbl.rows[i].cells[0], lic, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell(tbl.rows[i].cells[1], qs, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell(tbl.rows[i].cells[2], score, align=WD_ALIGN_PARAGRAPH.CENTER)
    style_cell(tbl.rows[i].cells[3], time, align=WD_ALIGN_PARAGRAPH.CENTER)

add_spacer()

# Note box
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = tbl.rows[0].cells[0]
set_cell_shading(cell, 'FFFBEB')
cell.text = ''
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
run = p.add_run('    Nota: ')
run.bold = True
run.font.color.rgb = ORANGE
run.font.size = Pt(10)
run = p.add_run(
    'En la Clase B, 3 preguntas valen 2 puntos cada una, por lo que el puntaje '
    'máximo es de 38 puntos (no 35). El puntaje de aprobación es 33/38.'
)
run.font.color.rgb = DARK_GRAY
run.font.size = Pt(10)

doc.add_page_break()

# ============================================================
#           13. SEGURIDAD
# ============================================================

add_heading_styled('13. Seguridad y protección de datos', level=1)
add_thin_line()

add_body(
    'La seguridad ha sido una prioridad en cada decisión de diseño y desarrollo '
    'de la plataforma:'
)

security = [
    ('Encriptación de contraseñas',
     'Todas las contraseñas se almacenan utilizando bcrypt, un algoritmo de encriptación '
     'de grado industrial. Las contraseñas nunca se guardan en texto plano ni en formatos '
     'reversibles.'),
    ('Autenticación por tokens',
     'Las sesiones de usuario se gestionan mediante tokens de autenticación seguros que '
     'se almacenan de forma protegida en el dispositivo del usuario.'),
    ('Base de datos segura',
     'PostgreSQL es una base de datos de nivel empresarial con encriptación de conexiones '
     'y respaldos automáticos gestionados por la infraestructura de alojamiento.'),
    ('Validación de datos',
     'Todas las entradas de usuario se validan tanto en el cliente como en el servidor, '
     'previniendo inyecciones de código y datos malformados.'),
    ('Verificación de pagos del lado del servidor',
     'Las compras se verifican directamente con los proveedores de pago (Mercado Pago '
     'y RevenueCat) desde el servidor, impidiendo la manipulación de planes desde el cliente.'),
    ('Sistema de roles',
     'Tres niveles de acceso protegen las funcionalidades sensibles: usuario estándar '
     '(acceso a la app), administrador (gestión de usuarios y contenido) y super administrador '
     '(mantenimiento técnico, uso exclusivo de WebMakerChile).'),
    ('Política de privacidad y términos legales',
     'La aplicación incluye una política de privacidad completa y términos de servicio '
     'accesibles desde la sección Legal, cumpliendo con las normativas de Apple y Google.'),
]

for title, desc in security:
    add_bullet(title, desc)

doc.add_page_break()

# ============================================================
#           14. ESTADO DE ENTREGA
# ============================================================

add_heading_styled('14. Estado de entrega y próximos pasos', level=1)
add_thin_line()

add_heading_styled('Entregado y funcionando', level=2, color=DARK_GREEN)

ready = [
    'Aplicación web completamente funcional y publicada en producción.',
    'Sistema completo de registro, inicio de sesión y gestión de cuentas.',
    'Base de datos con 1.591 preguntas (incluidas 306 oficiales de CONASET).',
    'Motor de exámenes con modo simulación y modo aprendizaje para 6 tipos de licencia.',
    'Temario completo con 11 capítulos y 115 secciones de estudio.',
    'Historial de exámenes, sistema de favoritos y seguimiento de progreso por categoría.',
    'Panel de administrador con gestión completa de usuarios (CRUD, búsqueda, filtros, detalles).',
    'Módulo de gestión de preguntas (crear, editar, buscar, filtrar, desactivar).',
    'Sistema de pagos con Mercado Pago integrado y operativo (web y Android).',
    'Integración con RevenueCat configurada para compras en iOS.',
    'Lectura en voz alta con inteligencia artificial (OpenAI TTS HD, voz Nova).',
    'Mascota copiloto con 5 estados visuales y animaciones interactivas.',
    'Efectos de sonido con variaciones aleatorias.',
    '68 imágenes ilustrativas con asignación automática por palabras clave.',
    'Barajado aleatorio de opciones para prevenir memorización por posición.',
    'Política de privacidad, términos de servicio y documentación legal completa.',
    'Compliance completo para Apple App Store (encriptación, privacidad, restaurar compras).',
    'Panel de configuración de voz para el usuario.',
    'Sistema de insignias por rachas de respuestas correctas.',
]

for item in ready:
    add_check(item)

add_spacer()

add_heading_styled('Próximos pasos (configuración externa)', level=2, color=ORANGE)

add_body(
    'Los siguientes elementos requieren configuración en plataformas externas y serán '
    'gestionados por WebMakerChile según el acuerdo de soporte:'
)

pending = [
    ('Publicación en Apple App Store',
     'WebMakerChile cuenta con la cuenta de Apple Developer necesaria. Se realizará la '
     'compilación nativa del proyecto y el envío a Apple para su revisión y publicación.'),
    ('Publicación en Google Play Store',
     'Requiere una cuenta de Google Play Developer ($25 USD, pago único). WebMakerChile '
     'se encargará de la compilación y publicación.'),
    ('Vinculación de RevenueCat con App Store Connect',
     'Necesario para que las compras in-app funcionen en iOS. Se completará una vez '
     'que la aplicación esté registrada en App Store Connect.'),
]

for title, desc in pending:
    add_warning(title, desc)

doc.add_page_break()

# ============================================================
#           15. GARANTÍA Y SOPORTE
# ============================================================

add_heading_styled('15. Garantía y soporte técnico', level=1)
add_thin_line()

add_body(
    'WebMakerChile se compromete a brindar soporte técnico para garantizar el correcto '
    'funcionamiento de la plataforma. Para cualquier consulta, incidencia o solicitud, '
    'puede comunicarse a través de los siguientes canales:'
)

add_spacer(8)

contact_info = [
    ('Correo electrónico', 'webmakerchile@gmail.com'),
    ('WhatsApp', '+56 9 6251 1821'),
    ('Desarrollador', 'Lucas - WebMakerChile'),
    ('Sitio web', 'webmakerchile.com'),
]

tbl = doc.add_table(rows=len(contact_info) + 1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

style_cell(tbl.rows[0].cells[0], 'Canal de contacto', bold=True, color=WHITE, bg='1B8A2F', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
style_cell(tbl.rows[0].cells[1], 'Detalle', bold=True, color=WHITE, bg='1B8A2F', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

for i, (channel, detail) in enumerate(contact_info, 1):
    style_cell(tbl.rows[i].cells[0], channel, bold=True, color=DARK_GREEN, bg='F0FFF4', size=10)
    style_cell(tbl.rows[i].cells[1], detail, size=10)

add_spacer(16)

add_body(
    'Agradecemos nuevamente la confianza depositada en WebMakerChile para el desarrollo '
    'de este proyecto. Estamos comprometidos con su éxito y disponibles para acompañarle '
    'en cada etapa del crecimiento de Practiquemos.cl.'
)

add_spacer(30)

# Footer
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
if os.path.exists(logo_path):
    run = p.add_run()
    run.add_picture(logo_path, width=Inches(1.5))

add_spacer(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('© 2026 webmakerchile.com — Todos los derechos reservados')
run.font.size = Pt(9)
run.font.color.rgb = MEDIUM_GRAY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Desarrollado con dedicación por Lucas - WebMakerChile')
run.font.size = Pt(9)
run.font.color.rgb = GREEN
run.italic = True

doc.save('Documento_Entrega_Practiquemos.docx')
print('Documento generado exitosamente.')
